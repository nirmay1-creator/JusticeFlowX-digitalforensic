from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import Optional
from pydantic import BaseModel
import uuid
import json

from database import get_db
from models import JusticeGPTChatSession, JusticeGPTChatMessage, JusticeGPTFileAnalysis
from services.ai_service import get_ai_response, analyze_file_content
from services.file_analyzer import process_uploaded_file

router = APIRouter(
    prefix="/api/justicegpt",
    tags=["JusticeGPT"]
)

class ChatRequest(BaseModel):
    message: str
    mode: str = "general"
    chat_id: Optional[str] = None

@router.post("/chat")
async def chat_endpoint(request: ChatRequest, db: Session = Depends(get_db)):
    chat_id = request.chat_id
    
    # If no chat_id, create a new session
    if not chat_id:
        chat_id = str(uuid.uuid4())
        new_session = JusticeGPTChatSession(session_uuid=chat_id, mode=request.mode)
        db.add(new_session)
        db.commit()
        db.refresh(new_session)
    
    # Fetch history
    history_records = db.query(JusticeGPTChatMessage).filter(
        JusticeGPTChatMessage.session_uuid == chat_id
    ).order_by(JusticeGPTChatMessage.created_at.asc()).all()
    
    history = [{"role": msg.role, "content": msg.content} for msg in history_records]
    
    # Save user message
    user_msg = JusticeGPTChatMessage(session_uuid=chat_id, role="user", content=request.message)
    db.add(user_msg)
    db.commit()
    
    # Call AI Service
    ai_reply = get_ai_response(request.message, request.mode, history)
    
    # Save AI message
    ai_msg = JusticeGPTChatMessage(session_uuid=chat_id, role="model", content=ai_reply)
    db.add(ai_msg)
    db.commit()
    
    return {
        "chat_id": chat_id,
        "reply": ai_reply,
        "mode": request.mode
    }

@router.post("/upload")
async def upload_endpoint(
    file: UploadFile = File(...),
    query: str = Form(...),
    mode: str = Form("general"),
    chat_id: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    if not chat_id:
        chat_id = str(uuid.uuid4())
        new_session = JusticeGPTChatSession(session_uuid=chat_id, mode=mode)
        db.add(new_session)
        db.commit()
        
    file_bytes = await file.read()
    
    # Extract text/content
    processed_data = process_uploaded_file(file.filename, file_bytes, file.content_type)
    
    # Call AI with extracted content
    if processed_data.get("content") and not processed_data["content"].startswith("Error"):
        ai_reply = analyze_file_content(processed_data["content"], file.content_type, query, mode)
    else:
        # Fallback for images or unsupported files if we want to pass bytes directly to gemini vision later
        # For now, just pass the filename and error
        ai_reply = f"Could not parse file content. File type might be unsupported or an image. Internal status: {processed_data.get('content')}"
        
        # If the model is Gemini 2.5 Pro, we could theoretically send the file bytes directly.
        # But this is a good starting point.
    
    # Mock threat level for now, AI could return this in structured format ideally.
    threat_level = "UNKNOWN"
    if "high risk" in ai_reply.lower() or "malicious" in ai_reply.lower() or "vulnerability" in ai_reply.lower():
        threat_level = "HIGH"
    elif "low risk" in ai_reply.lower() or "clean" in ai_reply.lower():
        threat_level = "LOW"
    
    analysis_record = JusticeGPTFileAnalysis(
        session_uuid=chat_id,
        filename=file.filename,
        file_hash=processed_data["hash"],
        file_type=file.content_type,
        analysis_results={"threat_level": threat_level, "summary": ai_reply}
    )
    db.add(analysis_record)
    
    # Save user message about upload
    user_msg = JusticeGPTChatMessage(session_uuid=chat_id, role="user", content=f"Uploaded file: {file.filename}\nQuery: {query}")
    ai_msg = JusticeGPTChatMessage(session_uuid=chat_id, role="model", content=ai_reply)
    db.add(user_msg)
    db.add(ai_msg)
    
    db.commit()
    
    # Confidence score is hardcoded mock for the requested UI style, but could be derived from AI.
    return {
        "chat_id": chat_id,
        "findings": ai_reply,
        "details": {
            "filename": file.filename,
            "hash": processed_data["hash"],
            "type": file.content_type,
            "threat_level": threat_level
        },
        "confidence": 94
    }

@router.get("/history/{chat_id}")
async def get_history(chat_id: str, db: Session = Depends(get_db)):
    records = db.query(JusticeGPTChatMessage).filter(JusticeGPTChatMessage.session_uuid == chat_id).order_by(JusticeGPTChatMessage.created_at.asc()).all()
    return [{"role": msg.role, "content": msg.content, "time": msg.created_at} for msg in records]

from fastapi.responses import FileResponse, JSONResponse
import os
from docx import Document

@router.post("/export")
async def export_report(request: ChatRequest, format: str = "pdf", db: Session = Depends(get_db)):
    chat_id = request.chat_id
    if not chat_id:
        raise HTTPException(status_code=400, detail="chat_id is required")
        
    records = db.query(JusticeGPTChatMessage).filter(JusticeGPTChatMessage.session_uuid == chat_id).order_by(JusticeGPTChatMessage.created_at.asc()).all()
    if not records:
        raise HTTPException(status_code=404, detail="Chat history not found")
        
    # Generate content
    content = "JusticeFlowX - JusticeGPT Forensic Report\n"
    content += "=" * 50 + "\n\n"
    for r in records:
        content += f"[{r.role.upper()}]:\n{r.content}\n\n"
        
    if format == "json":
        data = [{"role": r.role, "content": r.content} for r in records]
        return JSONResponse(content=data)
        
    elif format == "docx":
        doc = Document()
        doc.add_heading('JusticeFlowX - JusticeGPT Forensic Report', 0)
        for r in records:
            doc.add_heading(r.role.upper(), level=2)
            doc.add_paragraph(r.content)
            
        filepath = f"temp_report_{chat_id}.docx"
        doc.save(filepath)
        return FileResponse(filepath, filename=f"JusticeGPT_Report_{chat_id}.docx")
        
    elif format == "pdf":
        # Simplified PDF generation using reportlab if available, else fallback to text
        try:
            from reportlab.pdfgen import canvas
            filepath = f"temp_report_{chat_id}.pdf"
            c = canvas.Canvas(filepath)
            c.drawString(100, 800, "JusticeFlowX - JusticeGPT Forensic Report")
            y = 780
            for r in records:
                # Basic text wrapping logic could go here, but for now simple draw
                lines = r.content.split('\n')
                c.drawString(100, y, f"[{r.role.upper()}]:")
                y -= 15
                for line in lines[:10]: # Limit for basic PDF to avoid off-page
                    if y < 50:
                        c.showPage()
                        y = 800
                    c.drawString(120, y, line[:80]) # Truncate long lines
                    y -= 15
                y -= 10
            c.save()
            return FileResponse(filepath, filename=f"JusticeGPT_Report_{chat_id}.pdf")
        except ImportError:
            # Fallback to TXT if reportlab not installed properly
            filepath = f"temp_report_{chat_id}.txt"
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            return FileResponse(filepath, filename=f"JusticeGPT_Report_{chat_id}.txt")
    else:
        raise HTTPException(status_code=400, detail="Invalid format")

